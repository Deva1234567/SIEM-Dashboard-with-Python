import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import plotly.io as pio
import os
import shutil
from datetime import datetime, timedelta, date
import json
from docx import Document
from docx.shared import Inches
from io import BytesIO
import tempfile

# Paths
ALERTS_CSV = "C:/Users/devan/Desktop/SIEM/siem_dashboard/data/wazuh_alerts.csv"
CSS_FILE = "C:/Users/devan/Desktop/SIEM/siem_dashboard/static/style.css"
JS_FILE = "C:/Users/devan/Desktop/SIEM/siem_dashboard/static/script.js"

# Custom JSON serializer for Pandas Timestamp and datetime.date
def json_serializable(obj):
    if isinstance(obj, (pd.Timestamp, datetime)):
        return obj.isoformat()
    if isinstance(obj, date):
        return obj.isoformat()
    if isinstance(obj, pd._libs.tslibs.nattype.NaTType):
        return None
    raise TypeError(f"Type {type(obj)} not serializable")

# Load CSS
def load_css():
    if os.path.exists(CSS_FILE):
        with open(CSS_FILE, "r") as f:
            st.markdown(f"<style>{f.read()}</style>", unsafe_allow_html=True)

# Load JavaScript
def load_js():
    if os.path.exists(JS_FILE):
        with open(JS_FILE, "r") as f:
            st.components.v1.html(f"<script>{f.read()}</script>", height=0)

# Load alerts
@st.cache_data
def load_alerts(_refresh_key=0):
    if os.path.exists(ALERTS_CSV):
        df = pd.read_csv(ALERTS_CSV)
        df['timestamp'] = pd.to_datetime(df['timestamp'])
        df['date'] = df['timestamp'].dt.date
        # Ensure level is numeric
        df['level'] = pd.to_numeric(df['level'], errors='coerce')
        # Create a severity category based on numeric level
        df['severity_category'] = df['level'].apply(
            lambda x: 'High' if x >= 7 else 'Medium' if x >= 4 else 'Low'
        )
        return df
    return pd.DataFrame(columns=['timestamp', 'rule_description', 'agent_name', 'level', 'date', 'severity_category'])

# Function to generate Word document
def generate_word_doc(filtered_df, fig_bar_severity, fig_bar_agent, fig_pie, fig_line):
    doc = Document()
    doc.add_heading('SIEM Dashboard Report', 0)

    # Create a temporary directory for chart images
    temp_dir = tempfile.mkdtemp()

    try:
        # Overview Section
        doc.add_heading('Overview', level=1)
        doc.add_paragraph('This section provides a quick snapshot of the alerts.')
        doc.add_paragraph(f"Total Alerts: {len(filtered_df)}")
        doc.add_paragraph(f"High Severity (≥7): {len(filtered_df[filtered_df['level_numeric'] >= 7])}")
        doc.add_paragraph(f"Unique Agents: {len(filtered_df['agent_name'].unique())}")

        # Summary Section
        doc.add_heading('Summary', level=1)
        doc.add_paragraph('This section summarizes the alerts with key statistics.')
        avg_severity = filtered_df['level_numeric'].mean()
        most_active = filtered_df['agent_name'].mode().iloc[0] if not filtered_df.empty else "N/A"
        top_rule = filtered_df['rule_description'].mode().iloc[0] if not filtered_df.empty else "N/A"
        doc.add_paragraph(f"Average Severity: {avg_severity:.2f}")
        doc.add_paragraph(f"Most Active Agent: {most_active}")
        doc.add_paragraph(f"Top Rule: {top_rule}")

        # Charts Section
        doc.add_heading('Charts', level=1)
        doc.add_paragraph('These charts visualize the alert data.')

        # Save charts as images in the temporary directory
        charts = [
            (fig_bar_severity, "Alerts by Severity", "severity_bar.png"),
            (fig_bar_agent, "Alerts by Agent", "agent_bar.png"),
            (fig_pie, "Severity Breakdown", "severity_pie.png"),
            (fig_line, "Alerts Over Time", "alert_trend.png")
        ]
        for fig, title, filename in charts:
            if fig:  # Ensure the figure exists
                doc.add_heading(title, level=2)
                # Define the full path for the image file
                image_path = os.path.join(temp_dir, filename)
                # Write the chart image to the file
                pio.write_image(fig, image_path, format='png')
                # Embed the image in the Word document
                doc.add_picture(image_path, width=Inches(6))

        # Alerts Table Section
        doc.add_heading('Alerts Table', level=1)
        doc.add_paragraph('This table lists all alerts with their details.')
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Timestamp'
        hdr_cells[1].text = 'Rule Description'
        hdr_cells[2].text = 'Agent Name'
        hdr_cells[3].text = 'Severity Level'
        
        # Format the table data (remove HTML tags from severity level)
        filtered_df['severity_level'] = filtered_df['level_numeric'].astype(str)
        filtered_df['rule_desc'] = filtered_df['rule_description'].apply(lambda x: str(x).replace('<span class="tooltip" title="', '').replace('">', ': ').replace('</span>', '').replace('...', ''))
        for _, row in filtered_df[['timestamp_display', 'rule_desc', 'agent_name', 'severity_level']].iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = str(row['timestamp_display'])
            row_cells[1].text = str(row['rule_desc'])
            row_cells[2].text = str(row['agent_name'])
            row_cells[3].text = str(row['severity_level'])

        # Save the document to a BytesIO buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    finally:
        # Clean up the temporary directory and its contents
        try:
            shutil.rmtree(temp_dir)
        except Exception as e:
            st.warning(f"Could not clean up temporary directory: {e}")

# Initialize session state
if 'sidebar_visible' not in st.session_state:
    st.session_state.sidebar_visible = True
if 'refresh_key' not in st.session_state:
    st.session_state.refresh_key = 0
if 'selected_alert' not in st.session_state:
    st.session_state.selected_alert = None

# Main app
st.set_page_config(page_title="SIEM Dashboard", layout="wide")
load_css()
load_js()

# Load data
wazuh_df = load_alerts(st.session_state.refresh_key)

# Define filter variables with defaults
date_range = None
severity = []
agents = []
keyword = ""

# Sidebar
with st.sidebar:
    st.title("SIEM Controls")
    if st.button("Toggle Sidebar"):
        st.session_state.sidebar_visible = not st.session_state.sidebar_visible
    if st.session_state.sidebar_visible:
        st.markdown("")
        if st.button("Refresh Data"):
            st.session_state.refresh_key += 1
            st.markdown("""
                <div class="toast">Data refreshed</div>
                <script>
                    setTimeout(() => {
                        document.querySelector('.toast').style.opacity = '0';
                    }, 3000);
                </script>
            """, unsafe_allow_html=True)
            # Reload data on refresh
            wazuh_df = load_alerts(st.session_state.refresh_key)
        if not wazuh_df.empty:
            min_date = wazuh_df['date'].min()
            max_date = wazuh_df['date'].max()
            date_range = st.date_input(
                "Date Range",
                value=(min_date, max_date),
                min_value=min_date,
                max_value=max_date
            )
            severity = st.multiselect(
                "Severity Levels",
                options=sorted(wazuh_df['level'].unique()),
                default=sorted(wazuh_df['level'].unique())
            )
            agents = st.multiselect(
                "Agents",
                options=sorted(wazuh_df['agent_name'].unique()),
                default=sorted(wazuh_df['agent_name'].unique())
            )
            keyword = st.text_input("Search Description")

# Main content (centered via CSS)
st.title("SIEM Dashboard")

# Introduction
st.markdown("""
### Welcome to the SIEM Dashboard!  
This tool helps you keep an eye on security alerts from your systems. Think of it like a report card for your computer's safety—it shows you how many alerts happened, how serious they are, and lets you dig into the details. You can use the tabs below to explore different sections, and the sidebar on the left to filter the data.
""")

# Load and filter data
if wazuh_df.empty:
    st.error("No alerts found. Ensure data/wazuh_alerts.csv exists.")
else:
    filtered_df = wazuh_df.copy()  # Create a copy to avoid modifying the original
    # Preserve the numeric level for comparisons
    filtered_df['level_numeric'] = filtered_df['level']
    if isinstance(date_range, tuple) and len(date_range) == 2:
        start_date, end_date = date_range
        filtered_df = filtered_df[
            (filtered_df['date'] >= start_date) & (filtered_df['date'] <= end_date)
        ]
    if severity:
        filtered_df = filtered_df[filtered_df['level_numeric'].isin(severity)]
    if agents:
        filtered_df = filtered_df[filtered_df['agent_name'].isin(agents)]
    if keyword:
        filtered_df = filtered_df[filtered_df['rule_description'].str.contains(keyword, case=False, na=False)]

    # Tabs for different sections
    tab1, tab2, tab3, tab4 = st.tabs(["Overview", "Summary", "Charts", "Alerts"])

    # Overview Tab: Metrics
    with tab1:
        st.markdown("""
        #### Overview Tab
        This section gives you a quick snapshot of your alerts. It shows the total number of alerts, how many are high-severity (really important ones), and how many different systems (called agents) reported alerts.
        """)
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Total Alerts", len(filtered_df))
        with col2:
            st.metric("High Severity (≥7)", len(filtered_df[filtered_df['level_numeric'] >= 7]))
        with col3:
            st.metric("Unique Agents", len(filtered_df['agent_name'].unique()))

    # Summary Tab: Statistical Summary
    with tab2:
        st.markdown("""
        #### Summary Tab
        Here, you get a summary of your alerts. It tells you the average severity of alerts (how serious they are on average), which system is reporting the most alerts, and the most common type of alert.
        """)
        col1, col2, col3 = st.columns(3)
        with col1:
            avg_severity = filtered_df['level_numeric'].mean()
            st.metric("Average Severity", f"{avg_severity:.2f}")
        with col2:
            most_active = filtered_df['agent_name'].mode().iloc[0] if not filtered_df.empty else "N/A"
            st.metric("Most Active Agent", most_active)
        with col3:
            top_rule = filtered_df['rule_description'].mode().iloc[0] if not filtered_df.empty else "N/A"
            st.metric("Top Rule", top_rule)

    # Charts Tab: Alert Analysis
    with tab3:
        st.markdown("""
        #### Charts Tab
        This section uses charts to help you see patterns in your alerts. There are bar charts showing alerts by severity and agent, a pie chart for severity categories, and a line chart to see how alerts change over time.
        """)
        fig_bar_severity = None
        fig_bar_agent = None
        fig_pie = None
        fig_line = None
        if not filtered_df.empty:
            # Bar Chart: Severity Distribution
            severity_counts = filtered_df.groupby('level_numeric').size().reset_index(name='count')
            fig_bar_severity = px.bar(
                severity_counts,
                x='level_numeric',
                y='count',
                title="Alerts by Severity",
                color='level_numeric',
                color_continuous_scale='Viridis'
            )
            fig_bar_severity.update_layout(
                plot_bgcolor='rgba(0,0,0,0)',
                paper_bgcolor='rgba(0,0,0,0)',
                font_color='#333'
            )
            st.plotly_chart(fig_bar_severity, use_container_width=True)

            # Bar Chart: Agent Activity
            agent_counts = filtered_df.groupby('agent_name').size().reset_index(name='count')
            fig_bar_agent = px.bar(
                agent_counts,
                x='agent_name',
                y='count',
                title="Alerts by Agent",
                color='agent_name'
            )
            fig_bar_agent.update_layout(
                plot_bgcolor='rgba(0,0,0,0)',
                paper_bgcolor='rgba(0,0,0,0)',
                font_color='#333'
            )
            st.plotly_chart(fig_bar_agent, use_container_width=True)

            # Pie Chart: Severity Category
            category_counts = filtered_df.groupby('severity_category').size().reset_index(name='count')
            fig_pie = px.pie(
                category_counts,
                names='severity_category',
                values='count',
                title="Severity Breakdown",
                color_discrete_map={'Low': '#22c55e', 'Medium': '#f97316', 'High': '#ef4444'}
            )
            fig_pie.update_layout(
                plot_bgcolor='rgba(0,0,0,0)',
                paper_bgcolor='rgba(0,0,0,0)',
                font_color='#333'
            )
            st.plotly_chart(fig_pie, use_container_width=True)

            # Line Chart: Alert Trends
            trend_df = filtered_df.groupby('date').size().reset_index(name='count')
            fig_line = px.line(trend_df, x='date', y='count', title="Alerts Over Time")
            fig_line.update_layout(
                plot_bgcolor='rgba(0,0,0,0)',
                paper_bgcolor='rgba(0,0,0,0)',
                font_color='#333'
            )
            st.plotly_chart(fig_line, use_container_width=True)

    # Alerts Tab: Alert Table
    with tab4:
        st.markdown("""
        #### Alerts Tab
        This table lists all your alerts in detail. You can see when each alert happened, what it’s about, which system reported it, how serious it is, and click 'View' to see more details in a pop-up.
        """)
        if not filtered_df.empty:
            def severity_badge(level):
                color = "red" if level >= 7 else "orange" if level >= 4 else "green"
                return f'<span class="severity-badge" style="background-color:{color}">{level}</span>'
            # Format timestamp for display
            filtered_df['timestamp_display'] = filtered_df['timestamp'].dt.strftime('%Y-%m-%d %H:%M:%S')
            # Apply HTML formatting only for display
            filtered_df['level'] = filtered_df['level_numeric'].apply(severity_badge)
            # Add tooltip for rule_description
            filtered_df['rule_description'] = filtered_df['rule_description'].apply(
                lambda x: f'<span class="tooltip" title="{x}">{x[:30] + "..." if len(x) > 30 else x}</span>'
            )
            filtered_df['details'] = filtered_df.apply(
                lambda row: f'<button class="details-btn" onclick="showModal(\'{json.dumps(row.to_dict(), default=json_serializable)}\')">View</button>',
                axis=1
            )
            table_html = filtered_df[['timestamp_display', 'rule_description', 'agent_name', 'level', 'details']].to_html(
                classes="alert-table", escape=False, index=False
            )
            st.markdown(table_html, unsafe_allow_html=True)

            # Download as Word button
            buffer = generate_word_doc(filtered_df, fig_bar_severity, fig_bar_agent, fig_pie, fig_line)
            st.download_button(
                label="Download as Word",
                data=buffer,
                file_name="SIEM_Dashboard_Report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

            # Modal for alert details
            st.markdown("""
                <div id="alert-modal" class="modal">
                    <div class="modal-content">
                        <span class="close-btn" onclick="closeModal()">×</span>
                        <h3>Alert Details</h3>
                        <div id="modal-body"></div>
                    </div>
                </div>
            """, unsafe_allow_html=True)

        else:
            st.warning("No alerts match the selected filters.")